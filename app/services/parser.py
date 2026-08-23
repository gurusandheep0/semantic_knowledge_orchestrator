from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.config import settings


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedPage:
    number: int
    text: str


@dataclass(frozen=True)
class ParsedDocument:
    file_name: str
    file_type: str
    pages: list[ParsedPage]
    size_kb: float

    @property
    def text(self) -> str:
        return "\n\n".join(page.text for page in self.pages)


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("The text encoding could not be detected.")


def parse_document(payload: bytes, file_name: str) -> ParsedDocument:
    if not payload:
        raise DocumentParseError("The uploaded document is empty.")
    if len(payload) > settings.max_file_size_mb * 1024 * 1024:
        raise DocumentParseError(f"Documents must be smaller than {settings.max_file_size_mb} MB.")

    suffix = Path(file_name).suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(io.BytesIO(payload))
            pages = [ParsedPage(number=index, text=(page.extract_text() or "").strip()) for index, page in enumerate(reader.pages, start=1)]
        elif suffix == ".docx":
            document = Document(io.BytesIO(payload))
            text = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
            pages = [ParsedPage(number=1, text=text)]
        elif suffix in {".txt", ".md"}:
            pages = [ParsedPage(number=1, text=_decode_text(payload).strip())]
        else:
            raise DocumentParseError("Supported formats are PDF, DOCX, TXT, and Markdown.")
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Unable to parse {file_name}: {exc}") from exc

    pages = [page for page in pages if page.text]
    if not pages or len(" ".join(page.text for page in pages)) < 40:
        raise DocumentParseError("The document does not contain enough extractable text.")

    return ParsedDocument(file_name=file_name, file_type=suffix.lstrip("."), pages=pages, size_kb=round(len(payload) / 1024, 1))


def parse_text(text: str, file_name: str = "document.txt") -> ParsedDocument:
    clean = text.strip()
    if len(clean) < 40:
        raise DocumentParseError("The document does not contain enough extractable text.")
    suffix = Path(file_name).suffix.lower().lstrip(".") or "txt"
    return ParsedDocument(
        file_name=file_name,
        file_type=suffix,
        pages=[ParsedPage(number=1, text=clean)],
        size_kb=round(len(clean.encode("utf-8")) / 1024, 1),
    )
